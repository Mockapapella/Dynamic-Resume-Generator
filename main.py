import json
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx2pdf import convert

try:
    with open("resume.json", "r", encoding="utf8") as myfile:
        data = myfile.read()
except FileNotFoundError:
    with open("demo.json", "r", encoding="utf8") as myfile:
        data = myfile.read()

resume_json = json.loads(data)

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

document.styles["Normal"].paragraph_format.line_spacing = 1
document.styles["Normal"].paragraph_format.keep_together = True
document.styles["Normal"].paragraph_format.space_before = Pt(1)
document.styles["Normal"].paragraph_format.space_after = Pt(1)


# document.styles["Heading 1"].paragraph_format.space_before = Pt(4)
document.styles["Heading 1"].paragraph_format.space_after = Pt(1)
document.styles["Heading 2"].paragraph_format.space_before = Pt(6)
document.styles["Heading 2"].paragraph_format.space_after = Pt(1)
document.styles["Heading 3"].paragraph_format.space_before = Pt(4)
document.styles["Heading 3"].paragraph_format.space_after = Pt(1)

document.add_paragraph().add_run(resume_json["name"]).font.size = Pt(24)
document.add_paragraph().add_run(resume_json["title"]).font.size = Pt(16)
document.add_paragraph(resume_json["location"])
document.add_paragraph(resume_json["email"])
document.add_paragraph(resume_json["cell number"])
document.add_paragraph(resume_json["personal website"])
document.add_paragraph(resume_json["linkedin"])
document.add_paragraph(resume_json["github"])
document.add_paragraph("")
document.add_paragraph(resume_json["description"])


for component in resume_json:
    if component == "experience":
        check_true = []
        for subcomponent in resume_json[component]:
            check_true.append(resume_json[component][subcomponent]["include"])
        if True in check_true:
            document.add_heading("Experience", level=1)
            for subcomponent in resume_json[component]:
                if resume_json[component][subcomponent]["include"] is True:
                    document.add_heading(resume_json[component][subcomponent]["title"], level=2)
                    document.add_paragraph(
                        resume_json[component][subcomponent]["duration"][0]
                        + " - "
                        + resume_json[component][subcomponent]["duration"][1]
                    )
                    for item in resume_json[component][subcomponent]:
                        if (
                            item != "include"
                            and resume_json[component][subcomponent][item] is not None
                            and not isinstance(resume_json[component][subcomponent][item], list)
                            and item != "title"
                            and item != "duration"
                        ):
                            paragraph = document.add_paragraph("")
                            paragraph.add_run(item.title() + ": ").bold = True
                            paragraph.add_run(resume_json[component][subcomponent][item])
                        elif isinstance(resume_json[component][subcomponent][item], list):
                            if item == "techstack":
                                document.add_heading("Tech Stack", level=3)
                                for tech_count, technology in enumerate(
                                    resume_json[component][subcomponent][item]
                                ):
                                    paragraph = document.add_paragraph(
                                        resume_json[component][subcomponent][item][tech_count],
                                        style="List Bullet",
                                    )
                                    paragraph.paragraph_format.left_indent = Inches(0.75)
    elif component == "education":
        check_true = []
        for subcomponent in resume_json[component]:
            check_true.append(resume_json[component][subcomponent]["include"])
        if True in check_true:
            document.add_heading("Education", level=1)
            for subcomponent in resume_json[component]:
                if resume_json[component][subcomponent]["include"] is True:
                    document.add_heading(resume_json[component][subcomponent]["school"], level=2)
                    document.add_paragraph(
                        resume_json[component][subcomponent]["duration"][0]
                        + " - "
                        + resume_json[component][subcomponent]["duration"][1]
                    )
                    for item in resume_json[component][subcomponent]:
                        if (
                            item != "include"
                            and resume_json[component][subcomponent][item] is not None
                            and not isinstance(resume_json[component][subcomponent][item], list)
                            and item != "school"
                            and item != "duration"
                        ):
                            paragraph = document.add_paragraph("")
                            paragraph.add_run(item.title() + ": ").bold = True
                            paragraph.add_run(resume_json[component][subcomponent][item])
                        elif isinstance(resume_json[component][subcomponent][item], list):
                            if item == "activities and societies":
                                document.add_heading("Activities and Societies", level=3)
                                for activity_count, activity in enumerate(
                                    resume_json[component][subcomponent][item]
                                ):
                                    paragraph = document.add_paragraph(
                                        resume_json[component][subcomponent][item][activity_count],
                                        style="List Bullet",
                                    )
                                    paragraph.paragraph_format.left_indent = Inches(0.75)
    elif component == "licenses and certifications":
        check_true = []
        for subcomponent in resume_json[component]:
            check_true.append(resume_json[component][subcomponent]["include"])
        if True in check_true:
            document.add_heading("Licenses and Certifictions", level=1)
            for subcomponent in resume_json[component]:
                if resume_json[component][subcomponent]["include"] is True:
                    document.add_heading(resume_json[component][subcomponent]["name"], level=2)
                    document.add_paragraph(resume_json[component][subcomponent]["issued on"])
                    for item in resume_json[component][subcomponent]:
                        if (
                            item != "include"
                            and resume_json[component][subcomponent][item] is not None
                            and not isinstance(resume_json[component][subcomponent][item], list)
                            and item != "name"
                            and item != "issued on"
                        ):
                            paragraph = document.add_paragraph("")
                            paragraph.add_run(item.title() + ": ").bold = True
                            paragraph.add_run(resume_json[component][subcomponent][item])
    elif component == "volunteer experience":
        check_true = []
        for subcomponent in resume_json[component]:
            check_true.append(resume_json[component][subcomponent]["include"])
        if True in check_true:
            document.add_heading("Volunteer Experience", level=1)
            for subcomponent in resume_json[component]:
                if resume_json[component][subcomponent]["include"] is True:
                    document.add_heading(
                        resume_json[component][subcomponent]["organization"], level=2
                    )
                    document.add_paragraph(
                        resume_json[component][subcomponent]["duration"][0]
                        + " - "
                        + resume_json[component][subcomponent]["duration"][1]
                    )
                    for item in resume_json[component][subcomponent]:
                        if (
                            item != "include"
                            and resume_json[component][subcomponent][item] is not None
                            and not isinstance(resume_json[component][subcomponent][item], list)
                            and item != "organization"
                            and item != "duration"
                        ):
                            paragraph = document.add_paragraph("")
                            paragraph.add_run(item.title() + ": ").bold = True
                            paragraph.add_run(resume_json[component][subcomponent][item])

    elif component == "accomplishments":
        check_true = []
        for subcomponent in resume_json[component]:
            for subcomponent_item in resume_json[component][subcomponent]:
                check_true.append(
                    resume_json[component][subcomponent][subcomponent_item]["include"]
                )
        if True in check_true:
            document.add_heading("Accomplishments", level=1)
            for subcomponent in resume_json[component]:
                if subcomponent == "projects":
                    check_true = []
                    for project_subcomponent in resume_json[component][subcomponent]:
                        check_true.append(
                            resume_json[component][subcomponent][project_subcomponent]["include"]
                        )
                    if True in check_true:
                        document.add_heading("Projects", level=2)
                        for project_subcomponent in resume_json[component][subcomponent]:
                            if (
                                resume_json[component][subcomponent][project_subcomponent][
                                    "include"
                                ]
                                is True
                            ):
                                document.add_heading(
                                    resume_json[component][subcomponent][project_subcomponent][
                                        "name"
                                    ],
                                    level=3,
                                )
                                document.add_paragraph(
                                    resume_json[component][subcomponent][project_subcomponent][
                                        "duration"
                                    ][0]
                                    + " - "
                                    + resume_json[component][subcomponent][project_subcomponent][
                                        "duration"
                                    ][1]
                                )
                                for item in resume_json[component][subcomponent][
                                    project_subcomponent
                                ]:
                                    if (
                                        item != "include"
                                        and resume_json[component][subcomponent][
                                            project_subcomponent
                                        ][item]
                                        is not None
                                        and not isinstance(
                                            resume_json[component][subcomponent][
                                                project_subcomponent
                                            ][item],
                                            list,
                                        )
                                        and item != "name"
                                        and item != "duration"
                                    ):
                                        paragraph = document.add_paragraph("")
                                        paragraph.add_run(item.title() + ": ").bold = True
                                        paragraph.add_run(
                                            resume_json[component][subcomponent][
                                                project_subcomponent
                                            ][item]
                                        )
                                    elif isinstance(
                                        resume_json[component][subcomponent][project_subcomponent][
                                            item
                                        ],
                                        list,
                                    ):
                                        if item == "techstack":
                                            document.add_heading("Tech Stack", level=4)
                                            for tech_count, technology in enumerate(
                                                resume_json[component][subcomponent][
                                                    project_subcomponent
                                                ][item]
                                            ):
                                                paragraph = document.add_paragraph(
                                                    resume_json[component][subcomponent][
                                                        project_subcomponent
                                                    ][item][tech_count],
                                                    style="List Bullet",
                                                )
                                                paragraph.paragraph_format.left_indent = Inches(
                                                    0.75
                                                )
                if subcomponent == "honor and award":
                    check_true = []
                    for project_subcomponent in resume_json[component][subcomponent]:
                        check_true.append(
                            resume_json[component][subcomponent][project_subcomponent]["include"]
                        )
                    if True in check_true:
                        document.add_heading("Honor and Award", level=2)
                        for honor_subcomponent in resume_json[component][subcomponent]:
                            if (
                                resume_json[component][subcomponent][honor_subcomponent]["include"]
                                is True
                            ):
                                document.add_heading(
                                    resume_json[component][subcomponent][honor_subcomponent][
                                        "title"
                                    ],
                                    level=3,
                                )
                                document.add_paragraph(
                                    resume_json[component][subcomponent][honor_subcomponent][
                                        "issued on"
                                    ]
                                )
                                for item in resume_json[component][subcomponent][
                                    honor_subcomponent
                                ]:
                                    if (
                                        item != "include"
                                        and resume_json[component][subcomponent][
                                            honor_subcomponent
                                        ][item]
                                        is not None
                                        and not isinstance(
                                            resume_json[component][subcomponent][
                                                honor_subcomponent
                                            ][item],
                                            list,
                                        )
                                        and item != "title"
                                        and item != "issued on"
                                    ):
                                        paragraph = document.add_paragraph("")
                                        paragraph.add_run(item.title() + ": ").bold = True
                                        paragraph.add_run(
                                            resume_json[component][subcomponent][
                                                honor_subcomponent
                                            ][item]
                                        )
                if subcomponent == "language":
                    check_true = []
                    for project_subcomponent in resume_json[component][subcomponent]:
                        check_true.append(
                            resume_json[component][subcomponent][project_subcomponent]["include"]
                        )
                    if True in check_true:
                        document.add_heading("Language", level=2)
                        for language_subcomponent in resume_json[component][subcomponent]:
                            if (
                                resume_json[component][subcomponent][language_subcomponent][
                                    "include"
                                ]
                                is True
                            ):
                                document.add_heading(
                                    resume_json[component][subcomponent][language_subcomponent][
                                        "language"
                                    ],
                                    level=3,
                                )
                                for item in resume_json[component][subcomponent][
                                    language_subcomponent
                                ]:
                                    if (
                                        item != "include"
                                        and resume_json[component][subcomponent][
                                            language_subcomponent
                                        ][item]
                                        is not None
                                        and not isinstance(
                                            resume_json[component][subcomponent][
                                                language_subcomponent
                                            ][item],
                                            list,
                                        )
                                        and item != "language"
                                    ):
                                        paragraph = document.add_paragraph("")
                                        paragraph.add_run(item.title() + ": ").bold = True
                                        paragraph.add_run(
                                            resume_json[component][subcomponent][
                                                language_subcomponent
                                            ][item]
                                        )


Path(
    "generated_applications/{}/{}".format(
        resume_json["meta"]["company"], resume_json["meta"]["job"]
    )
).mkdir(parents=True, exist_ok=True)
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
document.save(
    "generated_applications/{}/{}/{} Resume - {}.docx".format(
        resume_json["meta"]["company"],
        resume_json["meta"]["job"],
        resume_json["name"],
        resume_json["meta"]["job"],
    )
)
convert(
    "generated_applications/{}/{}/{} Resume - {}.docx".format(
        resume_json["meta"]["company"],
        resume_json["meta"]["job"],
        resume_json["name"],
        resume_json["meta"]["job"],
    )
)
# convert("demo.docx", "output.pdf")
